"""
Document loader module for RAG Finance Tracking System.

This module handles loading and processing of various document types
including PDF, DOCX, TXT, and Excel files.
"""

from pathlib import Path
from typing import List, Union
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter


class DocumentLoader:
    """Handles loading and processing of documents."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document loader.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def load_text_file(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            List of Document objects
        """
        file_path = Path(file_path)
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        metadata = {"source": str(file_path), "type": "txt"}
        chunks = self.text_splitter.split_text(text)
        return [Document(page_content=chunk, metadata=metadata) for chunk in chunks]
    
    def load_pdf_file(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of Document objects
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required to load PDF files. Install it with: pip install pypdf")
        
        file_path = Path(file_path)
        reader = PdfReader(str(file_path))
        
        documents = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            metadata = {
                "source": str(file_path),
                "page": page_num + 1,
                "type": "pdf"
            }
            chunks = self.text_splitter.split_text(text)
            documents.extend([
                Document(page_content=chunk, metadata=metadata)
                for chunk in chunks
            ])
        
        return documents
    
    def load_docx_file(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of Document objects
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required to load DOCX files. Install it with: pip install python-docx")
        
        file_path = Path(file_path)
        doc = DocxDocument(str(file_path))
        
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        metadata = {"source": str(file_path), "type": "docx"}
        chunks = self.text_splitter.split_text(text)
        
        return [Document(page_content=chunk, metadata=metadata) for chunk in chunks]
    
    def load_directory(self, directory_path: Union[str, Path]) -> List[Document]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory_path: Path to the directory
            
        Returns:
            List of Document objects
        """
        directory_path = Path(directory_path)
        documents = []
        
        supported_extensions = {
            '.txt': self.load_text_file,
            '.pdf': self.load_pdf_file,
            '.docx': self.load_docx_file,
        }
        
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    loader_func = supported_extensions[file_path.suffix.lower()]
                    docs = loader_func(file_path)
                    documents.extend(docs)
                    print(f"Loaded {len(docs)} chunks from {file_path.name}")
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")
        
        return documents
    
    def load_document(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a single document, automatically detecting the file type.
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of Document objects
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == '.txt':
            return self.load_text_file(file_path)
        elif extension == '.pdf':
            return self.load_pdf_file(file_path)
        elif extension == '.docx':
            return self.load_docx_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
